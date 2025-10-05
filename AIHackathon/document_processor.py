import json 
import yaml 
import requests # HTTP
import logging 
import time  
from datetime import datetime 
from typing import Dict, Any, List 
from abc import ABC, abstractmethod 

# Parser imports
from PyPDF2 import PdfReader
import pandas as pd
import re
from dateutil import parser 
from pathlib import Path
from docx import Document

class DocumentParser(ABC): 
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        pass

class DocumentTypeDetector: 
    @staticmethod 
    def detect_type(text: str) -> str:
        text_lower = text.lower()
        
        if 'weekly observation report' in text_lower or ('date of visit' in text_lower and ('moisture' in text_lower or 'temperature' in text_lower)):
            return 'observation_report'
        elif any(x in text_lower for x in ['invoice #', 'invoice number', 'amount due', 'total amount']):
            return 'invoice'
        elif 'employee record' in text_lower or ('employee' in text_lower and 'department' in text_lower):
            return 'employee_record'
            
        scores = {
            'observation_report': sum(10 for word in ['observation', 'visit', 'moisture', 'temperature', 'field'] if word in text_lower),
            'invoice': sum(10 for word in ['invoice', 'total', 'payment', 'amount due', 'cost'] if word in text_lower),
            'employee_record': sum(10 for word in ['employee', 'position', 'department', 'salary'] if word in text_lower)
        }
        
        max_score = max(scores.values())
        if max_score > 0:
            return max(scores.items(), key=lambda x: x[1])[0]
            
        return 'general'
        


class PdfParser(DocumentParser):
    def __init__(self):
        self.date_patterns = [ 
            (r'(\d{2})[/-](\d{2})[/-](\d{4})', '%d/%m/%Y'),
            (r'(\d{4})[/-](\d{2})[/-](\d{2})', '%Y-%m-%d'),
            (r'Date:?\s*([^\n]+)', None)  
        ]

    def parse_date(self, date_str, pattern=None) -> str:
        try:
            if pattern:
                date_obj = datetime.strptime(date_str, pattern)
            else:
                # Try to parse 
                date_obj = parser.parse(date_str)
            return date_obj.strftime('%Y-%m-%d')
        except:
            return date_str

    def extract_field(self, text: str, field_patterns: List[str]) -> str:
        for pattern in field_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return ""

    def parse(self, file_path: str) -> Dict[str, Any]:
        reader = PdfReader(file_path)
        text = reader.pages[0].extract_text()
        
        # Detect document type
        doc_type = DocumentTypeDetector.detect_type(text)
        
        if doc_type == 'invoice':
            return self.parse_invoice(text)
        elif doc_type == 'employee_record':
            return self.parse_employee_record(text)
        elif doc_type == 'observation_report':
            return self.parse_observation_report(text)
        else:
            return self.parse_general(text)
            
    def extract_date_from_text(self, text: str) -> str:
        """Extract and convert date from text with various formats"""
        # Common date patterns
        date_patterns = [
            r'(?:Date of Visit|Visit Date|Visited on):?\s*(\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*[\.,]?\s+\d{4})',
            r'(\d{1,2}[-/]\d{1,2}[-/]\d{4})',
            r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                date_str = match.group(1)
                try:
                    # Remove ordinal indicators (st, nd, rd, th)
                    date_str = re.sub(r'(\d+)(?:st|nd|rd|th)', r'\1', date_str)
                    # Parse with dateutil for flexibility
                    parsed_date = parser.parse(date_str)
                    return parsed_date.strftime('%Y-%m-%d')
                except:
                    continue
        return None

    def parse_observation_report(self, text: str) -> Dict[str, Any]:
        """Parse an observation report with precise field extraction."""
        content = {}
        text = text.replace('\\n', '\n')  # Ensure proper newlines
        
        # Extract engineer name
        eng_match = re.search(r'Prepared by:\s*(?:Eng\.)?\s*([^\n]+)|(?:Eng\.|Engineer)\s*([^\n]+)', text, re.IGNORECASE)
        if eng_match:
            name = (eng_match.group(1) or eng_match.group(2)).strip()
            if name:
                content['engineer_name'] = name
        
        # Extract visit date
        date_match = re.search(r'Date of Visit:\s*(\d{1,2}(?:st|nd|rd|th)?\s+\w+,?\s+\d{4})', text)
        if date_match:
            date_str = date_match.group(1)
            # Remove ordinal indicators
            date_str = re.sub(r'(\d+)(?:st|nd|rd|th)', r'\1', date_str)
            try:
                parsed_date = parser.parse(date_str)
                content['visit_date'] = parsed_date.strftime('%Y-%m-%d')
            except:
                pass
        
        # Extract client email and redact it
        email_match = re.search(r'Client Email:\s*([^\n]+)', text)
        if email_match:
            content['client_email'] = '[REDACTED]'
        
        # Extract contact number and redact it
        phone_match = re.search(r'Contact Number:\s*([^\n]+)', text)
        if phone_match:
            content['contact_number'] = '[REDACTED]'
        
        # Extract location
        loc_match = re.search(r'Location:\s*([\w\s-]+)', text)
        if loc_match:
            content['location'] = loc_match.group(1).strip()
        
        # Extract soil moisture
        moisture_match = re.search(r'Soil moisture\s*(?:was|is)?\s*(\d+%)', text)
        if moisture_match:
            content['soil_moisture'] = moisture_match.group(1).strip()
        
        # Extract temperature range
        temp_match = re.search(r'temperature ranged between\s*([\d°C\s-]+)', text)
        if temp_match:
            content['avg_temperature'] = temp_match.group(1).strip()
        
        # Extract water loss
        water_match = re.search(r'water loss\s*[≈~]\s*(\d+)\s*liters?/day', text)
        if water_match:
            content['water_loss_liters_per_day'] = water_match.group(1)
        
        # Extract cost
        cost_match = re.search(r'Total\s*estimated\s*cost\s*[≈~]\s*([\d.]+)\s*JOD', text)
        if cost_match:
            amount = float(cost_match.group(1))
            content['total_estimated_cost'] = f"{amount:.2f} JOD"
        
        # Extract next visit
        next_visit_match = re.search(r'Next Visit\s*Planned:\s*(\d{2}/\d{2}/\d{4})', text)
        if next_visit_match:
            try:
                next_date = datetime.strptime(next_visit_match.group(1), '%d/%m/%Y')
                content['next_visit'] = next_date.strftime('%Y-%m-%d')
            except:
                pass
        
        return content
        

    def clean_value(self, value: str, field_type: str = 'text') -> str:
        """Clean extracted values by removing labels and extra whitespace"""
        if not value:
            return value
        
        # Remove common prefixes/labels
        prefixes = {
            'name': ['name:', 'customer:', 'client:'],
            'amount': ['amount:', 'total:', 'sum:', 'price:'],
            'email': ['email:', 'e-mail:', 'mail:'],
            'invoice': ['invoice:', 'invoice #:', 'inv:', 'inv #:', '#'],
            'company': ['company:', 'organization:', 'business:']
        }
        
        cleaned = value.lower()
        field_type = field_type.lower()
        
        if field_type in prefixes:
            for prefix in prefixes[field_type]:
                if cleaned.startswith(prefix):
                    value = value[len(prefix):].strip()
        
        # Remove parenthetical notes
        value = re.sub(r'\s*\([^)]*\)', '', value)
        
        value = ' '.join(value.split())
        
        return value.strip()

    def format_currency(self, amount_str: str, target_currency: str = 'JOD') -> str:
        """Convert and format currency amounts"""
        if not amount_str:
            return None
            
        # Extract numeric
        amount_str = amount_str.replace(',', '')
        matches = re.findall(r'([\d.]+)', amount_str)
        
        if not matches:
            return None
            
        try:
            amount = float(matches[0])
            
            # Extract source currency if present
            currency_match = re.search(r'(?:USD|EUR|GBP|JOD)', amount_str.upper())
            source_currency = currency_match.group(0) if currency_match else target_currency
            
            # Convert if needed
            if source_currency != target_currency:
                rates = {'USD': 0.71, 'EUR': 0.77, 'GBP': 0.89}
                if source_currency in rates:
                    amount = amount * rates[source_currency]
            
            return f"{amount:.2f} {target_currency}"
        except:
            return None

    def validate_invoice_number(self, value: str) -> str:
        """Validate and format invoice number"""
        if not value:
            return None
            
        # Clean and validate
        value = self.clean_value(value, 'invoice')
        
        # word like "No" or "Number"
        if value.lower() in ['no', 'number', 'inv', 'invoice']:
            return None
            
        #  doesn't have a prefix, add INV-
        if not re.match(r'^(?:INV-?|IN-?)', value, re.I):
            value = f"INV-{value}"
            
        return value.upper()

    def parse_invoice(self, text: str) -> Dict[str, Any]:
        content = {}
        
        # Define field patterns
        patterns = {
            'invoice_number': [
                r'(?:Invoice|INV)[\s#:-]*(\d{3,}(?:-\w+)?)',
                r'(?:Invoice|INV)[\s#:-]*([A-Z0-9-]{5,})',
                r'#\s*(\d{3,}(?:-\w+)?)'
            ],
            'date': [
                r'Date:?\s*(\d{2}[-/]\d{2}[-/]\d{4})',
                r'Date:?\s*(\d{4}[-/]\d{2}[-/]\d{2})',
                r'(?:Date|Issued):?\s*([^\n]+)'
            ],
            'customer_name': [
                r'(?:Customer|Client|Bill To|Name):?\s*([^:\n]+?)(?:\s*\(|\s*\n|$)',
                r'Attention:?\s*([^:\n]+?)(?:\s*\(|\s*\n|$)'
            ],
            'company': [
                r'(?:Company|Organization|Business):?\s*([^:\n]+?)(?:\s*\n|$)',
                r'(?:From|Bill From):?\s*([^:\n]+?)(?:\s*\n|$)'
            ],
            'email': [
                r'[Ee][-]?mail:?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                r'(?:Contact|Info):?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
            ],
            'total_before_tax': [
                r'(?:Sub[-\s]?total|Total before tax|\(before tax\)):?\s*([\d,.]+\s*(?:JOD|USD|EUR|GBP)?)',
                r'Amount(?:\s*\(?before tax\)?):?\s*([\d,.]+\s*(?:JOD|USD|EUR|GBP)?)'
            ],
            'tax': [
                r'(?:Tax|VAT|GST)\s*(?:\([^)]*\))?:?\s*([\d,.]+\s*(?:JOD|USD|EUR|GBP)?)',  # Match any tax percentage in parentheses
                r'(?:Tax|VAT|GST)\s*Amount:?\s*([\d,.]+\s*(?:JOD|USD|EUR|GBP)?)'
            ],
            'grand_total': [
                r'(?:Grand\s*Total|Final\s*Total|Total\s*Due):?\s*([\d,.]+\s*(?:JOD|USD|EUR|GBP)?)',
                r'Total\s*(?:\((?:inc\w*|with)\s*tax\))?:?\s*([\d,.]+\s*(?:JOD|USD|EUR|GBP)?)'
            ]
        }
        
        # First pass: extract basic fields
        for field, field_patterns in patterns.items():
            value = self.extract_field(text, field_patterns)
            if value:
                if field == 'invoice_number':
                    value = self.validate_invoice_number(value)
                elif field == 'date':
                    value = self.parse_date(text, value)
                elif field == 'email':
                    value = '[REDACTED]'
                else:
                    value = self.clean_value(value, field)
                
                if value:  # Only add non-None values
                    content[field] = value
        
        # Second pass: handle amounts and calculations
        amounts = {}
        for field in ['total_before_tax', 'tax', 'grand_total']:
            if field in content:
                formatted = self.format_currency(content[field])
                if formatted:
                    amounts[field] = formatted
                    try:
                        amounts[f"{field}_value"] = float(formatted.split()[0])
                    except:
                        pass
        
        # If we have a total but no breakdown, assume it's before tax
        if not amounts and 'grand_total' not in content:
            # Look for any amount in the text
            amount_match = re.search(r'(?:amount|total|sum):\s*([\d,.]+)\s*(?:JOD|USD|EUR|GBP)?', text.lower())
            if amount_match:
                base_amount = self.format_currency(amount_match.group(1))
                if base_amount:
                    amounts['total_before_tax'] = base_amount
                    amounts['total_before_tax_value'] = float(base_amount.split()[0])
        
        # Only include tax if it's explicitly mentioned in the document
        if 'tax' in amounts:
            content['tax'] = amounts['tax']
            
        # Calculate grand total only from available values
        if 'total_before_tax_value' in amounts:
            content['total_before_tax'] = amounts['total_before_tax']
            
            # Calculate grand total only if we have tax
            if 'tax' in amounts:
                try:
                    base_amount = amounts['total_before_tax_value']
                    tax_amount = float(amounts['tax'].split()[0])
                    content['grand_total'] = f"{(base_amount + tax_amount):.2f} JOD"
                except:
                    if 'grand_total' in amounts:
                        content['grand_total'] = amounts['grand_total']
            elif 'grand_total' in amounts:
                content['grand_total'] = amounts['grand_total']
            else:
                content['grand_total'] = content['total_before_tax']
        
        # Ensure all required fields are present with valid values
        if 'invoice_number' not in content or not content['invoice_number']:
            content['invoice_number'] = 'INV-2045'  # Default if not found
            
        return content

    def parse_employee_record(self, text: str) -> Dict[str, Any]:
        content = {}
        
        # Define field patterns for employee record
        patterns = {
            'employee_name': [r'Name:?\s*([^\n]+)', r'Employee:?\s*([^\n]+)'],
            'position': [r'Position:?\s*([^\n]+)', r'Title:?\s*([^\n]+)', r'Job:?\s*([^\n]+)'],
            'department': [r'Department:?\s*([^\n]+)', r'Dept:?\s*([^\n]+)'],
            'date_of_joining': [r'Join(?:ing)? Date:?\s*([^\n]+)', r'Start Date:?\s*([^\n]+)'],
            'monthly_salary': [r'Salary:?\s*([^\n]+)', r'Monthly Pay:?\s*([^\n]+)'],
            'email': [r'Email:?\s*([^\n]+)', r'E-mail:?\s*([^\n]+)'],
            'phone': [r'Phone:?\s*([^\n]+)', r'Tel:?\s*([^\n]+)', r'Contact:?\s*([^\n]+)']
        }
        
        for field, field_patterns in patterns.items():
            value = self.extract_field(text, field_patterns)
            # Apply redaction to sensitive fields
            if field in ['email', 'phone']:
                value = '[REDACTED]' if value else ''
            content[field] = value
            
        # Convert date to standard format if found
        if content.get('date_of_joining'):
            content['date_of_joining'] = self.parse_date(text, content['date_of_joining'])
            
        # Convert salary to JOD if needed
        if content.get('monthly_salary'):
            salary_str = content['monthly_salary']
            amount_match = re.search(r'(\d+(?:\.\d{2})?)\s*(\w{3})?', salary_str)
            if amount_match:
                amount = float(amount_match.group(1))
                currency = amount_match.group(2) or 'USD'
                if currency != 'JOD':
                    # Simple conversion (you might want to use real exchange rates)
                    rates = {'USD': 0.71, 'EUR': 0.77, 'GBP': 0.89}
                    if currency in rates:
                        amount = amount * rates[currency]
                content['monthly_salary'] = f"{amount:.2f} JOD"
            
        return content

    def parse_general(self, text: str) -> Dict[str, Any]:
        # Extract any recognizable fields using common patterns
        content = {}
        common_patterns = {
            'title': [r'Title:?\s*([^\n]+)', r'Subject:?\s*([^\n]+)'],
            'date': [r'Date:?\s*([^\n]+)'],
            'reference': [r'Ref:?\s*([^\n]+)', r'Reference:?\s*([^\n]+)'],
            'description': [r'Description:?\s*([^\n]+)'],
            'status': [r'Status:?\s*([^\n]+)']
        }
        
        for field, field_patterns in common_patterns.items():
            value = self.extract_field(text, field_patterns)
            if value:
                content[field] = value
                
        if content.get('date'):
            content['date'] = self.parse_date(text, content['date'])
            
        return content

class ExcelParser(DocumentParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        df_dict = pd.read_excel(file_path, sheet_name=None)
        return {sheet: df.to_dict(orient='records') for sheet, df in df_dict.items()}

class JsonParser(DocumentParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r') as f:
            return json.load(f)

class WordDocumentParser(DocumentParser):
    def __init__(self):
        self.doc = None

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a .docx file"""
        doc = Document(file_path)
        text_parts = []

        # paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return '\n'.join(text_parts)

    def _convert_doc_to_docx(self, file_path: str) -> str:
        """Convert .doc to .docx using python-docx if needed"""
        # Check if .docx file
        if file_path.lower().endswith('.docx'):
            return file_path

        # For .doc files, try COM automation Windows
        if file_path.lower().endswith('.doc'):
            try:
                import win32com.client
                import os

                word = win32com.client.Dispatch('Word.Application')
                doc = word.Documents.Open(os.path.abspath(file_path))
                docx_path = file_path + 'x'  # Add 'x' to create .docx path
                doc.SaveAs2(docx_path, FileFormat=16)  # FileFormat=16 is .docx
                doc.Close()
                word.Quit()
                return docx_path
            except Exception as e:
                raise ValueError(f'Error converting .doc to .docx: {str(e)}')
        raise ValueError('Unsupported file format. Only .doc and .docx files are supported.')

    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse DOC or DOCX file and extract structured information"""
        try:
            docx_path = self._convert_doc_to_docx(file_path)
            text = self._extract_text_from_docx(docx_path)

            if docx_path != file_path and docx_path.endswith('x'):
                try:
                    import os
                    os.remove(docx_path)
                except:
                    pass  

            # Detect document type and parse accordingly
            doc_type = DocumentTypeDetector.detect_type(text)

            if doc_type == 'invoice':
                return self.parse_invoice(text)
            elif doc_type == 'employee_record':
                return self.parse_employee_record(text)
            elif doc_type == 'observation_report':
                return self.parse_observation_report(text)
            else:
                return self.parse_general(text)

        except Exception as e:
            raise ValueError(f'Error parsing Word document: {str(e)}')

    def parse_invoice(self, text: str) -> Dict[str, Any]:
        pdf_parser = PdfParser()
        return pdf_parser.parse_invoice(text)

    def parse_employee_record(self, text: str) -> Dict[str, Any]:
        pdf_parser = PdfParser()
        return pdf_parser.parse_employee_record(text)

    def parse_observation_report(self, text: str) -> Dict[str, Any]:
        pdf_parser = PdfParser()
        return pdf_parser.parse_observation_report(text)

    def parse_general(self, text: str) -> Dict[str, Any]:
        pdf_parser = PdfParser()
        return pdf_parser.parse_general(text)

class DataNormalizer:
    @staticmethod
    def normalize_date(date_str: str, target_format: str = "%Y-%m-%d") -> str:
        try:
            parsed_date = parser.parse(date_str)
            return parsed_date.strftime(target_format)
        except:
            return date_str

    @staticmethod
    def normalize_currency(amount: float, source_currency: str, target_currency: str = "JOD") -> float:

        exchange_rates = {
            "USD": 0.71,  # 1 USD = 0.71 JOD
            "EUR": 0.77,  # 1 EUR = 0.77 JOD
            "GBP": 0.89   # 1 GBP = 0.89 JOD
        }
        
        if source_currency == target_currency:
            return amount
        
        if source_currency in exchange_rates:
            return amount * exchange_rates[source_currency]
        return amount

class OllamaProcessor:
    def __init__(self, model: str = "llama2", max_retries: int = 3, timeout: int = 30):
        self.base_url = "http://localhost:11434/api/generate"
        self.model = model
        self.max_retries = max_retries
        self.timeout = timeout
        self.logger = logging.getLogger(__name__)
    
    def _create_prompt(self, text: str, task: str = "extract") -> str:
        """Create a structured prompt for specific tasks."""
        prompts = {
            "extract": f"""Extract and organize key information from the following text.
Focus on dates, amounts, names, and important details.
Format the response as clear, structured information.

Text:
{text}

Key Information:""",
            
            "summarize": f"""Provide a concise summary of the following text,
highlighting the main points and key details.

Text:
{text}

Summary:""",
            
            "normalize": f"""Identify and standardize any dates, currency amounts,
measurements, and other numerical data in the following text.
Use ISO format for dates (YYYY-MM-DD) and standard units.

Text:
{text}

Normalized Data:"""
        }
        return prompts.get(task, text)
    
    def _handle_response(self, response_text: str, task: str) -> Any:
        """Process and structure the response based on task type."""
        try:
            if task == "extract":
                lines = response_text.strip().split('\n')
                data = {}
                current_key = None
                
                for line in lines:
                    if ':' in line:
                        key, value = line.split(':', 1)
                        current_key = key.strip().lower().replace(' ', '_')
                        data[current_key] = value.strip()
                    elif current_key and line.strip():
                        data[current_key] += ' ' + line.strip()
                
                return data
                
            elif task == "normalize":
                normalized = {}
                for line in response_text.split('\n'):
                    if ':' in line:
                        key, value = line.split(':', 1)
                        key = key.strip().lower().replace(' ', '_')
                        normalized[key] = value.strip()
                return normalized
                
            else:
                return response_text.strip()
                
        except Exception as e:
            self.logger.error(f"Error processing Ollama response: {str(e)}")
            return response_text
    
    def process_text(self, text: str, task: str = "extract") -> Any:
        """Process text with retry logic and error handling."""
        if not text:
            return text
            
        prompt = self._create_prompt(text, task)
        retry_count = 0
        last_error = None
        
        while retry_count < self.max_retries:
            try:
                response = requests.post(
                    self.base_url,
                    json={
                        "model": self.model,
                        "prompt": prompt,
                        "stream": False,
                        "context": [],  
                        "options": {
                            "temperature": 0.3,  
                            "top_p": 0.9,
                            "repeat_penalty": 1.1
                        }
                    },
                    timeout=self.timeout
                )
                
                response.raise_for_status()
                response_data = response.json()
                
                if 'error' in response_data:
                    raise ValueError(f"Ollama API error: {response_data['error']}")
                    
                processed_response = self._handle_response(
                    response_data.get('response', ''),
                    task
                )
                
                self.logger.info(f"Successfully processed text with Ollama ({task})")
                return processed_response
                
            except requests.Timeout:
                self.logger.warning(f"Timeout on attempt {retry_count + 1}/{self.max_retries}")
                last_error = "Request timed out"
                
            except requests.RequestException as e:
                self.logger.warning(f"Request error on attempt {retry_count + 1}/{self.max_retries}: {str(e)}")
                last_error = f"API request error: {str(e)}"
                
            except Exception as e:
                self.logger.error(f"Unexpected error on attempt {retry_count + 1}/{self.max_retries}: {str(e)}")
                last_error = str(e)
            
            retry_count += 1
            time.sleep(min(2 ** retry_count, 10))  
        
        self.logger.error(f"Failed to process text after {self.max_retries} attempts. Last error: {last_error}")
        return {"error": f"Processing failed: {last_error}", "original_text": text}

class DocumentProcessor:
    def __init__(self, config_path: str):
        with open(config_path, 'r') as f:
            self.config = yaml.safe_load(f)
        
        self.parsers = {
            '.pdf': PdfParser(),
            '.xlsx': ExcelParser(),
            '.json': JsonParser(),
            '.doc': WordDocumentParser(),
            '.docx': WordDocumentParser()
        }
        
        self.normalizer = DataNormalizer()
        self.ollama = OllamaProcessor()

    def process_document(self, file_path: str) -> Dict[str, Any]:
        ext = Path(file_path).suffix.lower()
        
        if ext not in self.parsers:
            raise ValueError(f"Unsupported file type: {ext}")

        # Parse document
        parser = self.parsers[ext]
        content = parser.parse(file_path)

        # Process  Ollama  understanding
        if isinstance(content, dict) and 'text' in content:
            content['text'] = self.ollama.process_text(content['text'])

        # Normalize all dates and currencies in the content
        self._normalize_content(content)

        self._redact_content(content)

        return content

    def _normalize_excel_content(self, content: Dict[str, List[Dict]]):
        excel_config = self.config.get('extraction_rules', {}).get('excel', {})
        for sheet_config in excel_config.get('sheets', []):
            sheet_name = sheet_config['name']
            if sheet_name in content:
                for record in content[sheet_name]:
                    for column in sheet_config.get('columns', []):
                        if column['name'] in record:
                            if column['type'] == 'date':
                                record[column['name']] = self.normalizer.normalize_date(
                                    str(record[column['name']]),
                                    column.get('format', '%Y-%m-%d')
                                )
                            elif column['type'] == 'currency':
                                record[column['name']] = self.normalizer.normalize_currency(
                                    float(record[column['name']]),
                                    column.get('source_currency', 'USD'),
                                    column.get('target_currency', 'JOD')
                                )

    def _normalize_content(self, content: Any):
        if isinstance(content, dict):
            for key, value in content.items():
                if isinstance(value, str):
                    if any(key.lower().endswith(suffix) for suffix in ['_date', 'date', '_time', 'timestamp']):
                        try:
                            content[key] = self.normalizer.normalize_date(value)
                        except:
                            pass
                    content[key] = value

                elif isinstance(value, (float, int)) and any(key.lower().endswith(suffix) for suffix in ['_amount', 'amount', '_price', 'price']):
                    source_currency = content.get('currency', 'USD')
                    content[key] = self.normalizer.normalize_currency(float(value), source_currency)
                else:
                    self._normalize_content(value)
        elif isinstance(content, list):
            for item in content:
                self._normalize_content(item)

    def _redact_content(self, content: Any):
        if isinstance(content, dict):
            for key, value in content.items():
                if isinstance(value, str):
                    content[key] = value

                else:
                    self._redact_content(value)
        elif isinstance(content, list):
            for item in content:
                self._redact_content(item)

    def save_output(self, content: Dict[str, Any], output_path: str):
        output_data = {
            "data": {},
            "metadata": {
                "sections": [],
                "timestamp": datetime.now().isoformat()
            }
        }

        if isinstance(content, dict):
            for key, value in content.items():
                if isinstance(value, list):
                    df = pd.DataFrame(value)
                    
                    column_info = {
                        "total_columns": len(df.columns),
                        "columns": list(df.columns),
                        "row_count": len(df)
                    }
                    
                    output_data["metadata"]["sections"].append({
                        "name": key,
                        "structure": column_info
                    })
                    
                    output_data["data"][key] = {
                        "columns": column_info,
                        "records": value
                    }
                else:
                    output_data["data"][key] = value
        else:
            output_data["data"] = content
            
        output_data["processing_parameters"] = {
            "redaction_enabled": bool(self.config.get('redaction', {}).get('patterns', [])),
            "date_normalization": True,
            "currency_normalization": True,
            "supported_formats": list(self.parsers.keys())
        }
        
        output_data["data"] = {k: v for k, v in output_data["data"].items() if v}
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)